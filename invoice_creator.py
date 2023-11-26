from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import pandas as pd
import calendar
from pathlib import Path
from argparse import ArgumentParser

class Session:
    def __init__(self, day, date, pupil, hrs, earned):
        self.day = day
        self.date = date
        self.pupil = pupil
        self.hrs = hrs
        self.earned = earned

    def __str__(self):
        return 'Session(' + self.day[:3] + ', ' + self.date + ' - ' + self.pupil + '(' + str(self.hrs) + ') - £' + str(self.earned) + ')'
    
    def __repr__(self):
        return 'Session(' + self.day[:3] + ', ' + self.date + ' - ' + self.pupil + '(' + str(self.hrs) + ') - £' + str(self.earned) + ')'

def read_log(filepath):
    df = pd.read_csv(filepath)
    df = df[['Day', 'Date', 'Pupils']]

    df = df.loc[df['Pupils'].notnull(), ['Day', 'Date', 'Pupils']]
    
    return df

def sessions_creator(df):
    pupils = []
    days = [i for i in df['Day']]
    dates = [i for i in df['Date']]

    i = 0
    j = 0

    sessions = {}

    for pupil in df['Pupils']:
        if pupil.isalpha() == True:
            pupils.append(pupil)

            sessions[j] = Session(days[i], dates[i], pupil, 1, 20)
            j += 1
            i += 1

        else:
            n_pupils = pupil.count(',')

            for n in range(n_pupils):
                pupil1, pupil2 = pupil.split(',', 1)

                sessions[j] = Session(days[i], dates[i], pupil1, 1, round(20, 2))
                j += 1

                pupil = pupil2

            sessions[j] = Session(days[i], dates[i], pupil, 1, round(20, 2))
            j += 1
            i += 1

    for j in range(len(sessions)):
        if sessions[j].pupil.isalpha() != True:
            pupil, hrs = sessions[j].pupil.split('(')
            hrs = int(hrs.replace(')', ''))

            sessions[j].pupil = pupil
            sessions[j].hrs = hrs
            sessions[j].earned = round(hrs * 20, 2)

    return sessions

def week_sort(sessions):
    weeks = {}
    week = []
    j = 0

    for i in range(len(sessions)):
        week.append(sessions[i])

        if (sessions[i].day == 'Sunday') and ((i == len(sessions)-1) or (sessions[i].date != sessions[i+1].date)):
            weeks[j] = week
            week = []
            j += 1

    weeks[j] = week

    return weeks

def create_invoice(sessions, weeks):
    day, month, year = sessions[0].date.split('/')
    month_name = calendar.month_name[int(month)]

    month_total = 0
    for i in range(len(sessions)):
        month_total += sessions[i].earned

    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)

    n_rows = len(sessions)+len(weeks)+2
    t = doc.add_table(rows=n_rows, cols=5)
    t.style = 'Table Grid'

    A = t.cell(0, 0).add_paragraph("Date")
    A.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    B = t.cell(0, 1).add_paragraph("Pupil")
    B.alignment = WD_ALIGN_PARAGRAPH.CENTER

    C = t.cell(0, 2).add_paragraph("Hours Worked")
    C.alignment = WD_ALIGN_PARAGRAPH.CENTER

    D = t.cell(0, 3).add_paragraph("Amount Earned")
    D.alignment = WD_ALIGN_PARAGRAPH.CENTER

    E = t.cell(0, 4).add_paragraph("Total")
    E.alignment = WD_ALIGN_PARAGRAPH.CENTER

    i = 1

    for j in range(len(weeks)):
        a = t.cell(i, 0)
        b = t.cell(i, 1)
        c = t.cell(i, 2)
        d = t.cell(i, 3)
        e = t.cell(i, 4)

        A = a.merge(b)
        B = A.merge(c)
        C = B.merge(d)
        D = C.merge(e)

        D = D.add_paragraph("Week {}".format(j+1))
        D.alignment = WD_ALIGN_PARAGRAPH.CENTER

        i += 1

        week_total = 0

        for session in weeks[j]:            
            t.cell(i, 1).text = str(session.pupil)
            t.cell(i, 2).text = str(session.hrs)
            t.cell(i, 3).text = '£' + str(session.earned)

            if t.cell(i-1, 0).text == str(session.date):
                a = t.cell(i-1, 0)
                b = t.cell(i, 0)
                a.merge(b)

            else:
                t.cell(i, 0).text = str(session.date)

            week_total += session.earned

            i += 1

        for k in range(len(weeks[j])):
            a = t.cell(i-1, 4)
            b = t.cell(i-1-k, 4)
            a.merge(b)

        t.cell(i-1, 4).text = '£' + str(week_total)
        t.cell(i-1, 4).vertical_alignment = WD_ALIGN_VERTICAL.TOP    

    a = t.cell(n_rows-1, 0)
    b = t.cell(n_rows-1, 1)
    c = t.cell(n_rows-1, 2)
    d = t.cell(n_rows-1, 3)
    t.cell(n_rows-1, 4).text = '£' + str(month_total)
    t.cell(n_rows-1, 4).vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM

    A = a.merge(b)
    B = A.merge(c)
    C = B.merge(d)
    C = C.add_paragraph('Total for ' + month_name + ' ' + str(year))
    C.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.save('{0} {1} Invoice.docx'.format(month_name, year))
    print("Successfully created invoice '{0} {1} Invoice.docx'.".format(month_name, year))


if __name__ == "__main__":
    parser = ArgumentParser(description="Generate an invoice")
    parser.add_argument('path', help="The filepath to the work log csv file.")
    arguments = parser.parse_args()
    
    file_path = Path(arguments.path)
    session_list = read_log(file_path)

    sessions = sessions_creator(session_list)
    weeks = week_sort(sessions)

    create_invoice(sessions, weeks)